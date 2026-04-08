import os
import tempfile
import threading
import unittest
from types import SimpleNamespace

import httpx
from pypdf import PdfWriter

from app.document_models import CanonicalBlock, KnowledgeChunkRecord, KnowledgeSourceFile, UploadStatus
from app.integrations import TemporalRuntime
from app.main import create_app
from app.routes import upload_knowledge_file
from app.storage import FilesystemStorageAdapter


class TestFallbackTemporalRuntime(TemporalRuntime):
    def __init__(self) -> None:
        super().__init__(target_hostport="temporal:7233", namespace="default")
        self.mode = "fallback"
        self.reason = "test_runtime"


def build_client(database_url=None):
    tempdir = tempfile.mkdtemp(prefix="emata-upload-")
    resolved_database_url = database_url or f"sqlite:///{os.path.join(tempdir, 'test.db')}"
    app = create_app(
        database_url=resolved_database_url,
        temporal_runtime=TestFallbackTemporalRuntime(),
    )
    transport = httpx.ASGITransport(app=app)
    client = httpx.AsyncClient(transport=transport, base_url="http://testserver")
    return app, client, resolved_database_url


def build_client_with_transport(database_url=None, raise_app_exceptions=True):
    tempdir = tempfile.mkdtemp(prefix="emata-upload-")
    resolved_database_url = database_url or f"sqlite:///{os.path.join(tempdir, 'test.db')}"
    app = create_app(
        database_url=resolved_database_url,
        temporal_runtime=TestFallbackTemporalRuntime(),
    )
    transport = httpx.ASGITransport(
        app=app,
        raise_app_exceptions=raise_app_exceptions,
    )
    client = httpx.AsyncClient(transport=transport, base_url="http://testserver")
    return app, client, resolved_database_url


class UploadApiTestCase(unittest.IsolatedAsyncioTestCase):
    @staticmethod
    def _build_valid_pdf_bytes() -> bytes:
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        output = tempfile.SpooledTemporaryFile()
        writer.write(output)
        output.seek(0)
        payload = output.read()
        output.close()
        return payload

    def test_filesystem_storage_adapter_returns_real_storage_path(self) -> None:
        with tempfile.TemporaryDirectory() as tempdir:
            adapter = FilesystemStorageAdapter(base_dir=tempdir)

            path = adapter.put_bytes(
                "org-1/workspace-finance/file-1/policy.txt",
                b"abc",
                "text/plain",
            )

            self.assertTrue(
                path.endswith("org-1\\workspace-finance\\file-1\\policy.txt")
                or path.endswith("org-1/workspace-finance/file-1/policy.txt")
            )
            self.assertTrue(os.path.exists(path))

    def test_filesystem_storage_adapter_round_trips_existing_absolute_path(self) -> None:
        with tempfile.TemporaryDirectory() as tempdir:
            adapter = FilesystemStorageAdapter(base_dir=tempdir)
            stored_path = adapter.put_bytes(
                "org-1/workspace-finance/file-1/policy.txt",
                b"abc",
                "text/plain",
            )

            local_path = adapter.get_to_local_path(stored_path, tempdir)

            self.assertEqual(local_path, stored_path)
            self.assertTrue(os.path.exists(local_path))

    async def test_get_upload_status_returns_storage_metadata(self) -> None:
        app, client, _database_url = build_client()
        source_file = KnowledgeSourceFile(
            id="upload-1",
            organization_id="org-acme",
            workspace_id="workspace-finance",
            scope="workspace",
            filename="policy.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_type="docx",
            storage_path="knowledge-source-files/org-acme/workspace-finance/upload-1/policy.docx",
            status=UploadStatus.PENDING,
            created_at="2026-03-30T00:00:00Z",
        )
        app.state.container.store.source_files[source_file.id] = source_file
        app.state.container.store.save_source_file(source_file)
        chunk = KnowledgeChunkRecord(
            id="upload-1-chunk-0",
            source_file_id="upload-1",
            organization_id="org-acme",
            workspace_id="workspace-finance",
            scope="workspace",
            title="policy",
            content="第一章\n报销审批正文",
            block_type="paragraph",
            section_path=["第一章"],
            page_number=2,
            sheet_name=None,
            slide_number=None,
            chunk_index=0,
            token_count_estimate=12,
            metadata={"parser": "mineru"},
        )
        app.state.container.store.chunks[chunk.id] = chunk
        app.state.container.store.save_chunk(chunk)

        response = await client.get("/api/v1/knowledge/uploads/upload-1")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["id"], "upload-1")
        self.assertEqual(payload["status"], "PENDING")
        self.assertEqual(payload["storage_path"], source_file.storage_path)
        self.assertEqual(payload["created_at"], "2026-03-30T00:00:00Z")
        self.assertEqual(payload["chunk_count"], 1)
        self.assertEqual(
            payload["ingestion_summary"],
            {
                "parser_backend": "mineru",
                "page_start": 2,
                "page_end": 2,
                "section_samples": ["第一章"],
                "block_types": ["paragraph"],
            },
        )
        await client.aclose()

    async def test_upload_endpoint_accepts_text_file_and_creates_chunks(self) -> None:
        app, client, _database_url = build_client()

        response = await client.post(
            "/api/v1/knowledge/uploads",
            files={"file": ("policy.txt", "报销制度正文".encode("utf-8"), "text/plain")},
            data={"workspace_id": "workspace-finance", "scope": "workspace"},
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["status"], "COMPLETED")
        self.assertEqual(payload["source_type"], "txt")
        self.assertEqual(
            payload["ingestion_summary"],
            {
                "parser_backend": "TXT",
                "page_start": None,
                "page_end": None,
                "section_samples": [],
                "block_types": ["paragraph"],
            },
        )
        self.assertTrue(
            any(
                chunk.source_file_id == payload["id"]
                for chunk in app.state.container.store.chunks.values()
            )
        )
        await client.aclose()

    async def test_upload_endpoint_returns_403_for_workspace_out_of_scope(self) -> None:
        _app, client, _database_url = build_client()

        response = await client.post(
            "/api/v1/knowledge/uploads",
            files={"file": ("policy.txt", "报销制度正文".encode("utf-8"), "text/plain")},
            data={"workspace_id": "workspace-secret", "scope": "workspace"},
        )

        self.assertEqual(response.status_code, 403)
        self.assertEqual(response.json()["detail"], "workspace_access_denied")
        await client.aclose()

    async def test_upload_endpoint_uses_normalized_source_type_for_txt_suffix(self) -> None:
        _app, client, _database_url = build_client()

        response = await client.post(
            "/api/v1/knowledge/uploads",
            files={
                "file": (
                    "policy.txt",
                    "报销制度正文".encode("utf-8"),
                    "application/octet-stream",
                )
            },
            data={"workspace_id": "workspace-finance", "scope": "workspace"},
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["source_type"], "txt")
        await client.aclose()

    async def test_upload_endpoint_accepts_docx_and_keeps_heading_metadata(self) -> None:
        from docx import Document

        app, client, _database_url = build_client()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            path = handle.name

        try:
            document = Document()
            document.add_heading("财务制度", level=1)
            document.add_paragraph("报销审批需要部门负责人确认。")
            document.save(path)

            with open(path, "rb") as payload_handle:
                response = await client.post(
                    "/api/v1/knowledge/uploads",
                    files={
                        "file": (
                            "policy.docx",
                            payload_handle.read(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    },
                    data={"workspace_id": "workspace-finance", "scope": "workspace"},
                )
        finally:
            os.remove(path)

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["source_type"], "docx")
        stored_chunks = [
            chunk
            for chunk in app.state.container.store.chunks.values()
            if chunk.source_file_id == payload["id"]
        ]
        self.assertTrue(stored_chunks)
        self.assertTrue(any(chunk.section_path == ["财务制度"] for chunk in stored_chunks))
        self.assertTrue(any("报销审批需要部门负责人确认" in chunk.content for chunk in stored_chunks))
        await client.aclose()

    async def test_upload_endpoint_keeps_parser_backend_in_ingestion_summary(self) -> None:
        app, client, _database_url = build_client()
        app.state.container.document_parser.parse_file = lambda *_args, **_kwargs: [
            CanonicalBlock(
                block_type="heading",
                text="第一章 总则",
                section_path=["第一章 总则"],
                page_number=1,
                metadata={"source_type": "pdf", "parser": "mineru"},
            ),
            CanonicalBlock(
                block_type="paragraph",
                text="第一条 正文内容。",
                section_path=["第一章 总则"],
                page_number=2,
                metadata={"source_type": "pdf", "parser": "mineru"},
            ),
        ]

        response = await client.post(
            "/api/v1/knowledge/uploads",
            files={"file": ("policy.pdf", b"%PDF-1.4 fake", "application/pdf")},
            data={"workspace_id": "workspace-finance", "scope": "workspace"},
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(
            payload["ingestion_summary"],
            {
                "parser_backend": "mineru",
                "page_start": 1,
                "page_end": 2,
                "section_samples": ["第一章 总则"],
                "block_types": ["paragraph"],
            },
        )
        await client.aclose()

    async def test_upload_endpoint_marks_pdf_as_failed_when_mineru_is_missing(self) -> None:
        app, client, _database_url = build_client()
        app.state.container.document_parser.registry._parsers["pdf"].executable = "missing-mineru"

        response = await client.post(
            "/api/v1/knowledge/uploads",
            files={
                "file": (
                    "policy.pdf",
                    self._build_valid_pdf_bytes(),
                    "application/pdf",
                )
            },
            data={"workspace_id": "workspace-finance", "scope": "workspace"},
        )

        self.assertEqual(response.status_code, 503)
        self.assertEqual(response.json()["detail"], "mineru_executable_not_found")
        failed_uploads = [
            item
            for item in app.state.container.store.source_files.values()
            if item.filename == "policy.pdf"
        ]
        self.assertTrue(failed_uploads)
        self.assertEqual(failed_uploads[-1].status.value, "FAILED")
        self.assertEqual(failed_uploads[-1].error_code, "mineru_executable_not_found")
        await client.aclose()

    async def test_upload_endpoint_rejects_invalid_pdf_before_invoking_parser(self) -> None:
        _app, client, _database_url = build_client()

        response = await client.post(
            "/api/v1/knowledge/uploads",
            files={"file": ("policy.pdf", b"%PDF-1.4 fake", "application/pdf")},
            data={"workspace_id": "workspace-finance", "scope": "workspace"},
        )

        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.json()["detail"], "invalid_pdf_file")
        await client.aclose()

    async def test_upload_endpoint_returns_504_when_parser_times_out(self) -> None:
        app, client, _database_url = build_client()
        app.state.container.document_parser.parse_file = lambda *_args, **_kwargs: (_ for _ in ()).throw(
            ValueError("parse_timeout")
        )

        response = await client.post(
            "/api/v1/knowledge/uploads",
            files={"file": ("policy.pdf", b"%PDF-1.4 fake", "application/pdf")},
            data={"workspace_id": "workspace-finance", "scope": "workspace"},
        )

        self.assertEqual(response.status_code, 504)
        self.assertEqual(response.json()["detail"], "parse_timeout")
        await client.aclose()

    async def test_upload_endpoint_returns_499_when_upload_is_canceled(self) -> None:
        app, client, _database_url = build_client()

        def canceled_ingest(**_kwargs):
            raise ValueError("upload_canceled")

        app.state.container.ingest_uploaded_file = canceled_ingest

        response = await client.post(
            "/api/v1/knowledge/uploads",
            files={"file": ("policy.txt", b"policy body", "text/plain")},
            data={"workspace_id": "workspace-finance", "scope": "workspace"},
        )

        self.assertEqual(response.status_code, 499)
        self.assertEqual(response.json()["detail"], "upload_canceled")
        await client.aclose()

    async def test_upload_route_streams_file_to_local_path_before_ingest(self) -> None:
        payload = b"chunked-upload-payload"
        captured = {}

        class FakeUploadFile:
            def __init__(self):
                self.filename = "policy.pdf"
                self.content_type = "application/pdf"
                self._cursor = 0
                self.read_sizes = []

            async def read(self, size: int = -1):
                self.read_sizes.append(size)
                if size is None or size <= 0:
                    raise AssertionError("route should read uploads in fixed-size chunks")
                if self._cursor >= len(payload):
                    return b""
                chunk = payload[self._cursor : self._cursor + size]
                self._cursor += len(chunk)
                return chunk

            async def close(self):
                return None

        class FakeRequest:
            async def is_disconnected(self):
                return False

        class FakeContainer:
            def ingest_uploaded_file(
                self,
                *,
                user,
                workspace_id,
                scope,
                filename,
                content_type,
                local_source_path=None,
                file_bytes=None,
                cancel_event=None,
            ):
                captured["workspace_id"] = workspace_id
                captured["scope"] = scope
                captured["filename"] = filename
                captured["content_type"] = content_type
                captured["cancel_event"] = cancel_event
                captured["local_source_path"] = local_source_path
                captured["file_bytes"] = file_bytes
                with open(local_source_path, "rb") as handle:
                    captured["stored_payload"] = handle.read()
                return KnowledgeSourceFile(
                    id="upload-streamed",
                    organization_id=user.organization_id,
                    workspace_id=workspace_id,
                    scope=scope,
                    filename=filename,
                    mime_type=content_type,
                    source_type="pdf",
                    storage_path="knowledge-source-files/org-acme/workspace-finance/upload-streamed/policy.pdf",
                    status=UploadStatus.COMPLETED,
                )

            def get_chunk_count_for_source_file(self, source_file_id):
                return 1 if source_file_id == "upload-streamed" else 0

            def get_ingestion_summary_for_source_file(self, source_file_id):
                if source_file_id != "upload-streamed":
                    return None
                return {
                    "parser_backend": "mineru",
                    "page_start": 1,
                    "page_end": 1,
                    "section_samples": ["上传流测试"],
                    "block_types": ["paragraph"],
                }

        upload_file = FakeUploadFile()
        response = await upload_knowledge_file(
            request=FakeRequest(),
            workspace_id="workspace-finance",
            scope="workspace",
            file=upload_file,
            container=FakeContainer(),
            user=SimpleNamespace(organization_id="org-acme"),
        )

        self.assertEqual(response.status, "COMPLETED")
        self.assertEqual(captured["stored_payload"], payload)
        self.assertIsNone(captured["file_bytes"])
        self.assertTrue(captured["local_source_path"])
        self.assertTrue(all(size == 1024 * 1024 for size in upload_file.read_sizes))
        self.assertGreaterEqual(len(upload_file.read_sizes), 2)

    async def test_upload_endpoint_marks_source_failed_when_parser_raises_runtime_error(self) -> None:
        app, client, _database_url = build_client_with_transport(raise_app_exceptions=False)

        def broken_parse(*_args, **_kwargs):
            raise ModuleNotFoundError("docx parser missing")

        app.state.container.document_parser.parse_file = broken_parse
        response = await client.post(
            "/api/v1/knowledge/uploads",
            files={
                "file": (
                    "broken.docx",
                    b"fake-docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"workspace_id": "workspace-finance", "scope": "workspace"},
        )

        self.assertEqual(response.status_code, 500)
        failed_uploads = [
            item
            for item in app.state.container.store.source_files.values()
            if item.filename == "broken.docx"
        ]
        self.assertTrue(failed_uploads)
        self.assertEqual(failed_uploads[-1].status.value, "FAILED")
        self.assertEqual(
            failed_uploads[-1].error_code,
            "upload_processing_failed:ModuleNotFoundError",
        )
        self.assertIn("docx parser missing", failed_uploads[-1].error_message or "")
        await client.aclose()

    async def test_uploaded_chunks_remain_searchable_after_container_restart(self) -> None:
        with tempfile.TemporaryDirectory() as tempdir:
            database_url = f"sqlite:///{os.path.join(tempdir, 'test.db')}"
            first_app, client, _ = build_client(database_url=database_url)

            upload_response = await client.post(
                "/api/v1/knowledge/uploads",
                files={"file": ("policy.txt", "差旅审批必须走 ERP".encode("utf-8"), "text/plain")},
                data={"workspace_id": "workspace-finance", "scope": "workspace"},
            )
            self.assertEqual(upload_response.status_code, 201)
            await client.aclose()
            first_app.state.container.close()

            restarted_app, restarted_client, _ = build_client(database_url=database_url)
            search_response = await restarted_client.get(
                "/api/v1/knowledge/search",
                params={"workspace_id": "workspace-finance", "query": "差旅审批 ERP"},
            )

            self.assertEqual(search_response.status_code, 200)
            payload = search_response.json()
            self.assertTrue(any(item["title"] == "policy" for item in payload["items"]))
            await restarted_client.aclose()
            restarted_app.state.container.close()

    async def test_get_upload_status_returns_404_when_missing(self) -> None:
        _app, client, _database_url = build_client()

        response = await client.get("/api/v1/knowledge/uploads/missing-upload")

        self.assertEqual(response.status_code, 404)
        await client.aclose()

    async def test_get_upload_status_returns_403_when_out_of_scope(self) -> None:
        app, client, _database_url = build_client()
        source_file = KnowledgeSourceFile(
            id="upload-foreign",
            organization_id="org-other",
            workspace_id="workspace-finance",
            scope="workspace",
            filename="foreign.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_type="docx",
            storage_path="knowledge-source-files/org-other/workspace-finance/upload-foreign/foreign.docx",
            status=UploadStatus.PENDING,
            created_at="2026-03-30T00:00:00Z",
        )
        app.state.container.store.source_files[source_file.id] = source_file
        app.state.container.store.save_source_file(source_file)

        response = await client.get("/api/v1/knowledge/uploads/upload-foreign")

        self.assertEqual(response.status_code, 403)
        await client.aclose()

    async def test_list_uploads_returns_recent_history_with_chunk_counts(self) -> None:
        app, client, _database_url = build_client()
        latest = KnowledgeSourceFile(
            id="upload-latest",
            organization_id="org-acme",
            workspace_id="workspace-finance",
            scope="workspace",
            filename="latest.pdf",
            mime_type="application/pdf",
            source_type="pdf",
            storage_path="knowledge-source-files/org-acme/workspace-finance/upload-latest/latest.pdf",
            status=UploadStatus.COMPLETED,
            created_at="2026-03-30T01:00:00Z",
        )
        failed = KnowledgeSourceFile(
            id="upload-failed",
            organization_id="org-acme",
            workspace_id="workspace-finance",
            scope="workspace",
            filename="failed.pdf",
            mime_type="application/pdf",
            source_type="pdf",
            storage_path="knowledge-source-files/org-acme/workspace-finance/upload-failed/failed.pdf",
            status=UploadStatus.FAILED,
            error_code="parse_timeout",
            error_message="parse_timeout",
            created_at="2026-03-30T00:30:00Z",
        )
        shared = KnowledgeSourceFile(
            id="upload-shared",
            organization_id="org-acme",
            workspace_id=None,
            scope="shared",
            filename="shared.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_type="docx",
            storage_path="knowledge-source-files/org-acme/shared/upload-shared/shared.docx",
            status=UploadStatus.COMPLETED,
            created_at="2026-03-29T23:00:00Z",
        )
        app.state.container.store.source_files[latest.id] = latest
        app.state.container.store.source_files[failed.id] = failed
        app.state.container.store.source_files[shared.id] = shared
        app.state.container.store.save_source_file(latest)
        app.state.container.store.save_source_file(failed)
        app.state.container.store.save_source_file(shared)
        chunk = KnowledgeChunkRecord(
            id="upload-latest-chunk-0",
            source_file_id="upload-latest",
            organization_id="org-acme",
            workspace_id="workspace-finance",
            scope="workspace",
            title="latest",
            content="pdf chunk",
            block_type="paragraph",
            section_path=["Section"],
            page_number=1,
            sheet_name=None,
            slide_number=None,
            chunk_index=0,
            token_count_estimate=5,
            metadata={"parser": "mineru", "page_end": 2},
        )
        app.state.container.store.chunks[chunk.id] = chunk
        app.state.container.store.save_chunk(chunk)
        chunk_two = KnowledgeChunkRecord(
            id="upload-latest-chunk-1",
            source_file_id="upload-latest",
            organization_id="org-acme",
            workspace_id="workspace-finance",
            scope="workspace",
            title="latest",
            content="pdf chunk 2",
            block_type="table",
            section_path=["Section", "Appendix"],
            page_number=3,
            sheet_name=None,
            slide_number=None,
            chunk_index=1,
            token_count_estimate=6,
            metadata={"parser": "mineru", "page_end": 4},
        )
        app.state.container.store.chunks[chunk_two.id] = chunk_two
        app.state.container.store.save_chunk(chunk_two)

        response = await client.get(
            "/api/v1/knowledge/uploads",
            params={"workspace_id": "workspace-finance", "limit": 5},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual([item["id"] for item in payload["items"]], ["upload-latest", "upload-failed", "upload-shared"])
        self.assertEqual(payload["items"][0]["chunk_count"], 2)
        self.assertEqual(
            payload["items"][0]["ingestion_summary"],
            {
                "parser_backend": "mineru",
                "page_start": 1,
                "page_end": 4,
                "section_samples": ["Section", "Section / Appendix"],
                "block_types": ["paragraph", "table"],
            },
        )
        self.assertEqual(payload["items"][1]["error_code"], "parse_timeout")
        self.assertIsNone(payload["items"][1]["ingestion_summary"])
        self.assertEqual(payload["items"][2]["scope"], "shared")
        await client.aclose()

    def test_ingest_uploaded_file_rolls_back_saved_chunks_when_indexing_fails(self) -> None:
        app, client, _database_url = build_client()
        container = app.state.container
        container.document_parser.parse_file = lambda *_args, **_kwargs: [
            CanonicalBlock(block_type="paragraph", text="第一段。"),
            CanonicalBlock(block_type="heading", text="第二章", section_path=["第二章"]),
            CanonicalBlock(block_type="paragraph", text="第二段。"),
        ]
        calls = {"count": 0}

        def broken_index(chunk):
            calls["count"] += 1
            if calls["count"] == 2:
                raise RuntimeError("milvus down")

        container._index_chunk = broken_index

        with self.assertRaises(RuntimeError):
            container.ingest_uploaded_file(
                user=container.get_current_user(),
                workspace_id="workspace-finance",
                scope="workspace",
                filename="policy.txt",
                content_type="text/plain",
                file_bytes="第一段。第二段。".encode("utf-8"),
            )

        failed_upload = next(
            item for item in container.store.source_files.values() if item.filename == "policy.txt"
        )
        self.assertEqual(failed_upload.status.value, "FAILED")
        self.assertEqual(container.get_chunk_count_for_source_file(failed_upload.id), 0)
        app.state.container.close()

    def test_ingest_uploaded_file_marks_canceled_and_keeps_zero_chunks(self) -> None:
        app, client, _database_url = build_client()
        container = app.state.container
        container.document_parser.parse_file = lambda *_args, **_kwargs: [
            CanonicalBlock(block_type="paragraph", text="第一段。"),
            CanonicalBlock(block_type="paragraph", text="第二段。"),
        ]
        cancel_event = threading.Event()
        cancel_event.set()

        source_file = container.ingest_uploaded_file(
            user=container.get_current_user(),
            workspace_id="workspace-finance",
            scope="workspace",
            filename="policy.txt",
            content_type="text/plain",
            file_bytes="第一段。第二段。".encode("utf-8"),
            cancel_event=cancel_event,
        )

        self.assertEqual(source_file.status.value, "CANCELED")
        self.assertEqual(source_file.error_code, "upload_canceled")
        self.assertEqual(container.get_chunk_count_for_source_file(source_file.id), 0)
        app.state.container.close()


if __name__ == "__main__":
    unittest.main()
